import sys
import os
import time
import docx
from pii_redactor import PiiRedactor

def redact_paragraph(paragraph, redactor):
    """
    Redacts PII in a single paragraph while preserving run-level formatting.
    """
    text = paragraph.text
    if not text.strip():
        return
        
    matches = redactor.analyze_text(text)
    if not matches:
        return
        
    for m in matches:
        start_idx = m["start"]
        end_idx = m["end"]
        fake_text = m["fake"]
        
        # Re-build character-to-run mapping
        char_mapping = []
        for r_idx, run in enumerate(paragraph.runs):
            for char_idx in range(len(run.text)):
                char_mapping.append((r_idx, char_idx))
                
        # If mapping fails or is out of sync, fall back to merging all runs and replacing
        if len(char_mapping) != len(paragraph.text) or start_idx >= len(char_mapping) or (end_idx - 1) >= len(char_mapping):
            if paragraph.runs:
                redacted_full = paragraph.text[:start_idx] + fake_text + paragraph.text[end_idx:]
                paragraph.runs[0].text = redacted_full
                for run in paragraph.runs[1:]:
                    run.text = ""
            continue
            
        start_run_idx, start_char_in_run = char_mapping[start_idx]
        end_run_idx, end_char_in_run = char_mapping[end_idx - 1]
        
        if start_run_idx == end_run_idx:
            # Match is within a single run
            run = paragraph.runs[start_run_idx]
            run.text = run.text[:start_char_in_run] + fake_text + run.text[end_char_in_run + 1:]
        else:
            # Match spans multiple runs
            run_start = paragraph.runs[start_run_idx]
            run_start.text = run_start.text[:start_char_in_run] + fake_text
            
            # Clear text in intermediate runs
            for r_idx in range(start_run_idx + 1, end_run_idx):
                paragraph.runs[r_idx].text = ""
                
            run_end = paragraph.runs[end_run_idx]
            run_end.text = run_end.text[end_char_in_run + 1:]

def redact_cell(cell, redactor):
    """
    Redacts a table cell, including its paragraphs and any nested tables.
    """
    for p in cell.paragraphs:
        redact_paragraph(p, redactor)
    for table in cell.tables:
        redact_table(table, redactor)

def redact_table(table, redactor):
    """
    Redacts a table.
    """
    for row in table.rows:
        for cell in row.cells:
            redact_cell(cell, redactor)

def redact_header_footer(header_footer, redactor):
    """
    Redacts all paragraphs and tables in a header or footer.
    """
    if not header_footer:
        return
    for p in header_footer.paragraphs:
        redact_paragraph(p, redactor)
    for t in header_footer.tables:
        redact_table(t, redactor)

def redact_docx(input_path, output_path):
    print(f"Loading document: {input_path}")
    t0 = time.time()
    doc = docx.Document(input_path)
    print(f"Document loaded in {time.time()-t0:.2f}s")
    
    print("Initializing PII Redactor (loading spaCy and Presidio, may take up to 45 seconds)...")
    t0 = time.time()
    redactor = PiiRedactor()
    print(f"Redactor initialized in {time.time()-t0:.2f}s")
    
    # 1. Redact body paragraphs
    print("Redacting body paragraphs...")
    t0 = time.time()
    for i, p in enumerate(doc.paragraphs):
        redact_paragraph(p, redactor)
        if (i+1) % 100 == 0:
            print(f"  Processed {i+1}/{len(doc.paragraphs)} paragraphs...")
    print(f"Body paragraphs redacted in {time.time()-t0:.2f}s")
    
    # 2. Redact tables
    print("Redacting tables (including nested tables)...")
    t0 = time.time()
    for i, t in enumerate(doc.tables):
        redact_table(t, redactor)
        if (i+1) % 10 == 0:
            print(f"  Processed {i+1}/{len(doc.tables)} tables...")
    print(f"Tables redacted in {time.time()-t0:.2f}s")
    
    # 3. Redact headers and footers
    print("Redacting headers and footers...")
    t0 = time.time()
    for i, section in enumerate(doc.sections):
        redact_header_footer(section.header, redactor)
        redact_header_footer(section.footer, redactor)
        
        # Check first page and even page headers/footers
        if hasattr(section, 'first_page_header') and section.first_page_header:
            redact_header_footer(section.first_page_header, redactor)
        if hasattr(section, 'first_page_footer') and section.first_page_footer:
            redact_header_footer(section.first_page_footer, redactor)
        if hasattr(section, 'even_page_header') and section.even_page_header:
            redact_header_footer(section.even_page_header, redactor)
        if hasattr(section, 'even_page_footer') and section.even_page_footer:
            redact_header_footer(section.even_page_footer, redactor)
            
    print(f"Headers and footers redacted in {time.time()-t0:.2f}s")
    
    # Save the output
    print(f"Saving redacted document to: {output_path}")
    doc.save(output_path)
    print("Document saved successfully!")

if __name__ == "__main__":
    input_file = "Red Herring Prospectus.docx"
    output_file = "Red Herring Prospectus_Redacted.docx"
    
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
        
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
        
    redact_docx(input_file, output_file)
